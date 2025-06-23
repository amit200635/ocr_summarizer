import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from transformers import pipeline
from keybert import KeyBERT
from docx import Document
import os
import cv2
import numpy as np
import docx
import re
import unicodedata

# Set Tesseract path for Windows
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Preprocess image for better OCR
def preprocess(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    return gray

# Clean extracted text
def clean_text(text):
    text = unicodedata.normalize("NFKD", text)
    text = re.sub(r'[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD]', '', text)
    return text

# Extract text from supported files
def extract_text(file_path):
    ext = file_path.split('.')[-1].lower()
    text = ""

    if ext == "pdf":
        pages = convert_from_path(file_path)
        for i, page in enumerate(pages):
            img = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            img = preprocess(img)
            text += f"\n--- Page {i + 1} ---\n"
            text += pytesseract.image_to_string(img)

    elif ext in ["jpg", "jpeg", "png"]:
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)

    elif ext == "docx":
        doc = docx.Document(file_path)
        text = '\n'.join([para.text for para in doc.paragraphs])

    elif ext == "txt":
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()

    else:
        raise ValueError("❌ Unsupported file type.")

    return text

# Flask-callable function
def process_file(file_path):
    print(f"📂 Processing file: {file_path}")
    text = extract_text(file_path)
    cleaned_text = clean_text(text)

    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
    summary = summarizer(cleaned_text[:1024])[0]['summary_text']

    kw_model = KeyBERT()
    keywords = kw_model.extract_keywords(cleaned_text, top_n=10)
    keywords = [kw[0] for kw in keywords]
    cleaned_keywords = ', '.join(keywords)

    doc = Document()
    doc.add_heading('Extracted OCR/Text', level=1)
    for para in cleaned_text.split('\n'):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)

    doc.add_heading('Keywords', level=1)
    doc.add_paragraph(cleaned_keywords)

    output_file = os.path.join(os.getcwd(), "ocr_summary_output.docx")
    doc.save(output_file)

    return output_file, summary, cleaned_keywords


def process_file(file_path):
    print(f"📂 File selected: {file_path}")
    print("🔍 Extracting text...")
    text = extract_text(file_path)

    print("🧹 Cleaning text...")
    cleaned_text = clean_text(text)

    print("📄 Summarizing...")
    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
    summary = summarizer(cleaned_text[:1024])[0]['summary_text']

    print("🔑 Extracting keywords...")
    kw_model = KeyBERT()
    keywords = kw_model.extract_keywords(cleaned_text, top_n=10)
    keywords = [kw[0] for kw in keywords]
    cleaned_keywords = ', '.join(keywords)

    print("📝 Creating Word document...")
    doc = Document()
    doc.add_heading('Extracted OCR/Text', level=1)
    for para in cleaned_text.split('\n'):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)

    doc.add_heading('Keywords', level=1)
    doc.add_paragraph(cleaned_keywords)

    output_file = os.path.join("static", "ocr_summary_output.docx")
    doc.save(output_file)
    print(f"✅ Word document saved as '{output_file}' in {os.getcwd()}")

    return output_file, summary, cleaned_keywords
